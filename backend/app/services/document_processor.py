"""
Document processing service for parsing and chunking files.
Supports PDF, DOCX, TXT, and CSV formats.
"""
import asyncio
import io
from typing import List, Dict, Any, Tuple
from pathlib import Path
import PyPDF2
from docx import Document as DocxDocument
import pandas as pd
from app.config import settings
from app.core.logging import get_logger
from app.models.enums import FileType

logger = get_logger(__name__)


class DocumentChunk:
    """Represents a chunk of document text with metadata."""
    
    def __init__(
        self,
        text: str,
        chunk_index: int,
        metadata: Dict[str, Any]
    ):
        self.text = text
        self.chunk_index = chunk_index
        self.metadata = metadata


class DocumentProcessor:
    """Service for processing and chunking documents."""
    
    def __init__(self):
        """Initialize document processor."""
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP
    
    async def process_file(
        self,
        file_content: bytes,
        filename: str,
        file_type: FileType
    ) -> Tuple[List[DocumentChunk], Dict[str, Any]]:
        """
        Process a file and return chunks with metadata.
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            file_type: Type of file
        
        Returns:
            Tuple of (chunks, file_metadata)
        """
        logger.info(f"Processing file: {filename} (type: {file_type})")
        
        # Extract text based on file type
        if file_type == FileType.PDF:
            text, metadata = await self._process_pdf(file_content, filename)
        elif file_type == FileType.DOCX:
            text, metadata = await self._process_docx(file_content, filename)
        elif file_type == FileType.TXT:
            text, metadata = await self._process_txt(file_content, filename)
        elif file_type == FileType.CSV:
            text, metadata = await self._process_csv(file_content, filename)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Chunk the text
        chunks = self._chunk_text(text, filename, metadata)
        
        logger.info(f"Created {len(chunks)} chunks from {filename}")
        
        return chunks, metadata
    
    async def _process_pdf(self, file_content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Process PDF file."""
        loop = asyncio.get_event_loop()
        
        def extract_pdf():
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                text_parts.append(f"[Page {page_num + 1}]\n{text}")
            
            metadata = {
                "pages": len(pdf_reader.pages),
                "has_metadata": len(pdf_reader.metadata) > 0 if pdf_reader.metadata else False
            }
            
            return "\n\n".join(text_parts), metadata
        
        return await loop.run_in_executor(None, extract_pdf)
    
    async def _process_docx(self, file_content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Process DOCX file."""
        loop = asyncio.get_event_loop()
        
        def extract_docx():
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "sections": len(doc.sections)
            }
            
            return "\n\n".join(text_parts), metadata
        
        return await loop.run_in_executor(None, extract_docx)
    
    async def _process_txt(self, file_content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Process TXT file."""
        text = file_content.decode('utf-8', errors='ignore')
        
        metadata = {
            "lines": len(text.split('\n')),
            "characters": len(text)
        }
        
        return text, metadata
    
    async def _process_csv(self, file_content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Process CSV file."""
        loop = asyncio.get_event_loop()
        
        def extract_csv():
            csv_file = io.BytesIO(file_content)
            df = pd.read_csv(csv_file)
            
            # Convert DataFrame to text representation
            text_parts = [f"CSV Data from {filename}:\n"]
            text_parts.append(f"Columns: {', '.join(df.columns)}\n")
            
            # Add each row as text
            for idx, row in df.iterrows():
                row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
                text_parts.append(f"Row {idx + 1}: {row_text}")
            
            metadata = {
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": list(df.columns)
            }
            
            return "\n".join(text_parts), metadata
        
        return await loop.run_in_executor(None, extract_csv)
    
    def _chunk_text(
        self,
        text: str,
        filename: str,
        file_metadata: Dict[str, Any]
    ) -> List[DocumentChunk]:
        """
        Chunk text with overlap.
        
        Args:
            text: Full document text
            filename: Original filename
            file_metadata: Metadata from file processing
        
        Returns:
            List of DocumentChunk objects
        """
        # Simple word-based chunking with overlap
        words = text.split()
        chunks = []
        
        start = 0
        chunk_index = 0
        
        while start < len(words):
            # Get chunk words
            end = start + self.chunk_size
            chunk_words = words[start:end]
            chunk_text = " ".join(chunk_words)
            
            # Create chunk with metadata
            chunk = DocumentChunk(
                text=chunk_text,
                chunk_index=chunk_index,
                metadata={
                    "filename": filename,
                    "chunk_index": chunk_index,
                    "start_word": start,
                    "end_word": min(end, len(words)),
                    "total_words": len(words),
                    **file_metadata
                }
            )
            
            chunks.append(chunk)
            
            # Move start position with overlap
            start += self.chunk_size - self.chunk_overlap
            chunk_index += 1
        
        return chunks


# Global document processor instance
_document_processor: DocumentProcessor = None


def get_document_processor() -> DocumentProcessor:
    """
    Get the global document processor instance.
    
    Returns:
        DocumentProcessor: Document processor
    """
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessor()
    return _document_processor
